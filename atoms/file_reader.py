"""
Atomic component for reading different file types.
Handles PDF, Word, Excel, CSV, and text files.
"""

import io
from pathlib import Path
from typing import Union, BinaryIO, List, Dict, Any, Optional
from abc import ABC, abstractmethod
import pandas as pd

# PDF processing
import pdfplumber
from pypdf import PdfReader

# Word processing
from docx import Document as DocxDocument

# Excel processing
import openpyxl

from config.logging_config import logger
from core.exceptions import UnsupportedFileTypeError, DocumentProcessingError
from core.types import DocumentType


class BaseFileReader(ABC):
    """Abstract base class for file readers."""
    
    @abstractmethod
    def read(self, file_path: Union[str, Path, BinaryIO]) -> Dict[str, Any]:
        """Read file and return structured content."""
        pass
    
    @abstractmethod
    def extract_tables(self, file_path: Union[str, Path, BinaryIO]) -> List[Dict[str, Any]]:
        """Extract tables from the file."""
        pass


class PDFReader(BaseFileReader):
    """Reader for PDF files with table extraction."""
    
    def read(self, file_path: Union[str, Path, BinaryIO]) -> Dict[str, Any]:
        """
        Read PDF file and extract text and metadata.
        
        Args:
            file_path: Path to PDF file or file-like object
            
        Returns:
            Dictionary with text content, pages, and metadata
        """
        logger.info(f"Reading PDF file: {file_path}")
        
        try:
            result = {
                "pages": [],
                "metadata": {},
                "tables": []
            }
            
            # Extract text using pdfplumber (better for tables)
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text() or ""
                    result["pages"].append({
                        "page_number": page_num,
                        "text": page_text,
                        "width": page.width,
                        "height": page.height
                    })
                    
                    # Extract tables from this page
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if table and len(table) > 1:
                            result["tables"].append({
                                "page_number": page_num,
                                "table_index": table_idx,
                                "headers": table[0] if table else [],
                                "rows": table[1:] if len(table) > 1 else []
                            })
            
            # Extract metadata using pypdf
            with open(file_path, "rb") as f:
                pdf_reader = PdfReader(f)
                if pdf_reader.metadata:
                    result["metadata"] = {
                        "title": pdf_reader.metadata.get("/Title", ""),
                        "author": pdf_reader.metadata.get("/Author", ""),
                        "subject": pdf_reader.metadata.get("/Subject", ""),
                        "creator": pdf_reader.metadata.get("/Creator", ""),
                        "page_count": len(pdf_reader.pages)
                    }
            
            logger.info(f"Successfully read PDF with {len(result['pages'])} pages")
            return result
            
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise DocumentProcessingError(f"Failed to read PDF: {e}")
    
    def extract_tables(self, file_path: Union[str, Path, BinaryIO]) -> List[Dict[str, Any]]:
        """Extract all tables from PDF."""
        result = self.read(file_path)
        return result.get("tables", [])


class WordReader(BaseFileReader):
    """Reader for Word documents (.docx)."""
    
    def read(self, file_path: Union[str, Path, BinaryIO]) -> Dict[str, Any]:
        """
        Read Word document and extract content.
        
        Args:
            file_path: Path to Word file
            
        Returns:
            Dictionary with paragraphs, tables, and metadata
        """
        logger.info(f"Reading Word file: {file_path}")
        
        try:
            doc = DocxDocument(file_path)
            
            result = {
                "paragraphs": [],
                "tables": [],
                "metadata": {}
            }
            
            # Extract paragraphs
            for para_idx, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    result["paragraphs"].append({
                        "index": para_idx,
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal"
                    })
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    "table_index": table_idx,
                    "headers": [],
                    "rows": []
                }
                
                for row_idx, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    if row_idx == 0:
                        table_data["headers"] = row_data
                    else:
                        table_data["rows"].append(row_data)
                
                result["tables"].append(table_data)
            
            # Extract metadata
            core_props = doc.core_properties
            result["metadata"] = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else ""
            }
            
            logger.info(f"Successfully read Word doc with {len(result['paragraphs'])} paragraphs")
            return result
            
        except Exception as e:
            logger.error(f"Error reading Word document: {e}")
            raise DocumentProcessingError(f"Failed to read Word document: {e}")
    
    def extract_tables(self, file_path: Union[str, Path, BinaryIO]) -> List[Dict[str, Any]]:
        """Extract all tables from Word document."""
        result = self.read(file_path)
        return result.get("tables", [])


class ExcelReader(BaseFileReader):
    """Reader for Excel files (.xlsx, .xls)."""
    
    def read(self, file_path: Union[str, Path, BinaryIO]) -> Dict[str, Any]:
        """
        Read Excel file and extract all sheets.
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            Dictionary with sheets data and metadata
        """
        logger.info(f"Reading Excel file: {file_path}")
        
        try:
            result = {
                "sheets": [],
                "metadata": {}
            }
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # Clean column names
                df.columns = [str(col).strip() for col in df.columns]
                
                sheet_data = {
                    "sheet_name": sheet_name,
                    "headers": df.columns.tolist(),
                    "rows": df.values.tolist(),
                    "row_count": len(df),
                    "column_count": len(df.columns),
                    "summary": self._generate_summary(df)
                }
                
                result["sheets"].append(sheet_data)
            
            result["metadata"] = {
                "sheet_count": len(excel_file.sheet_names),
                "sheet_names": excel_file.sheet_names
            }
            
            logger.info(f"Successfully read Excel with {len(result['sheets'])} sheets")
            return result
            
        except Exception as e:
            logger.error(f"Error reading Excel file: {e}")
            raise DocumentProcessingError(f"Failed to read Excel file: {e}")
    
    def _generate_summary(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Generate summary statistics for numeric columns."""
        summary = {}
        
        for col in df.select_dtypes(include=['number']).columns:
            summary[col] = {
                "mean": float(df[col].mean()) if not pd.isna(df[col].mean()) else None,
                "sum": float(df[col].sum()) if not pd.isna(df[col].sum()) else None,
                "min": float(df[col].min()) if not pd.isna(df[col].min()) else None,
                "max": float(df[col].max()) if not pd.isna(df[col].max()) else None
            }
        
        return summary
    
    def extract_tables(self, file_path: Union[str, Path, BinaryIO]) -> List[Dict[str, Any]]:
        """Extract all sheets as tables."""
        result = self.read(file_path)
        return result.get("sheets", [])


class CSVReader(BaseFileReader):
    """Reader for CSV files."""
    
    def read(
        self, 
        file_path: Union[str, Path, BinaryIO],
        encoding: str = "utf-8",
        delimiter: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Read CSV file.
        
        Args:
            file_path: Path to CSV file
            encoding: File encoding
            delimiter: Custom delimiter (auto-detected if None)
            
        Returns:
            Dictionary with table data and metadata
        """
        logger.info(f"Reading CSV file: {file_path}")
        
        try:
            # Auto-detect delimiter if not provided
            if delimiter is None:
                import csv
                with open(file_path, 'r', encoding=encoding) as f:
                    sample = f.read(8192)
                    sniffer = csv.Sniffer()
                    try:
                        dialect = sniffer.sniff(sample)
                        delimiter = dialect.delimiter
                    except csv.Error:
                        delimiter = ','
            
            df = pd.read_csv(file_path, encoding=encoding, delimiter=delimiter)
            
            # Clean column names
            df.columns = [str(col).strip() for col in df.columns]
            
            result = {
                "headers": df.columns.tolist(),
                "rows": df.values.tolist(),
                "row_count": len(df),
                "column_count": len(df.columns),
                "summary": self._generate_summary(df),
                "metadata": {
                    "encoding": encoding,
                    "delimiter": delimiter
                }
            }
            
            logger.info(f"Successfully read CSV with {result['row_count']} rows")
            return result
            
        except Exception as e:
            logger.error(f"Error reading CSV file: {e}")
            raise DocumentProcessingError(f"Failed to read CSV file: {e}")
    
    def _generate_summary(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Generate summary statistics for numeric columns."""
        summary = {}
        
        for col in df.select_dtypes(include=['number']).columns:
            summary[col] = {
                "mean": float(df[col].mean()) if not pd.isna(df[col].mean()) else None,
                "sum": float(df[col].sum()) if not pd.isna(df[col].sum()) else None,
                "min": float(df[col].min()) if not pd.isna(df[col].min()) else None,
                "max": float(df[col].max()) if not pd.isna(df[col].max()) else None
            }
        
        return summary
    
    def extract_tables(self, file_path: Union[str, Path, BinaryIO]) -> List[Dict[str, Any]]:
        """Return CSV as a single table."""
        result = self.read(file_path)
        return [{
            "headers": result["headers"],
            "rows": result["rows"]
        }]


class TextReader(BaseFileReader):
    """Reader for plain text files."""
    
    def read(
        self, 
        file_path: Union[str, Path, BinaryIO],
        encoding: str = "utf-8"
    ) -> Dict[str, Any]:
        """
        Read plain text file.
        
        Args:
            file_path: Path to text file
            encoding: File encoding
            
        Returns:
            Dictionary with text content
        """
        logger.info(f"Reading text file: {file_path}")
        
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            
            result = {
                "text": content,
                "line_count": content.count('\n') + 1,
                "char_count": len(content),
                "word_count": len(content.split()),
                "metadata": {
                    "encoding": encoding
                }
            }
            
            logger.info(f"Successfully read text file with {result['word_count']} words")
            return result
            
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            raise DocumentProcessingError(f"Failed to read text file: {e}")
    
    def extract_tables(self, file_path: Union[str, Path, BinaryIO]) -> List[Dict[str, Any]]:
        """Text files typically don't have tables."""
        return []


class FileReaderFactory:
    """Factory for creating appropriate file readers."""
    
    _readers = {
        DocumentType.PDF: PDFReader,
        DocumentType.WORD: WordReader,
        DocumentType.EXCEL: ExcelReader,
        DocumentType.CSV: CSVReader,
        DocumentType.TEXT: TextReader
    }
    
    @classmethod
    def get_reader(cls, file_type: DocumentType) -> BaseFileReader:
        """
        Get appropriate reader for file type.
        
        Args:
            file_type: Type of document
            
        Returns:
            Appropriate file reader instance
        """
        reader_class = cls._readers.get(file_type)
        if not reader_class:
            raise UnsupportedFileTypeError(file_type.value)
        return reader_class()
    
    @classmethod
    def get_reader_by_extension(cls, extension: str) -> BaseFileReader:
        """
        Get reader based on file extension.
        
        Args:
            extension: File extension (with or without dot)
            
        Returns:
            Appropriate file reader instance
        """
        ext = extension.lower().lstrip('.')
        type_mapping = {
            'pdf': DocumentType.PDF,
            'docx': DocumentType.WORD,
            'doc': DocumentType.WORD,
            'xlsx': DocumentType.EXCEL,
            'xls': DocumentType.EXCEL,
            'csv': DocumentType.CSV,
            'txt': DocumentType.TEXT
        }
        
        file_type = type_mapping.get(ext)
        if not file_type:
            raise UnsupportedFileTypeError(ext)
        
        return cls.get_reader(file_type)